#!/usr/bin/env python3
"""Minimal but faithful Markdown -> .docx converter for the Aegis-40 FER master.

Handles the constructs actually used in Aegis40-FER-master.md:
  headings (#..####), pipe tables (with :---: alignment), bold (**), inline code (`),
  bullet/numbered lists, blockquotes (>), fenced code blocks (```), horizontal rules (---).

Usage:  py -3 scripts/md_to_docx.py <input.md> <output.docx>
"""
import io
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

BODY_FONT = "Arial"
MONO_FONT = "Consolas"
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text, base_bold=False, size=None, color=None):
    """Add a run sequence honouring **bold** and `code`."""
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(tok)
        if base_bold:
            run.bold = True
        if size is not None and run.font.size is None:
            run.font.size = size
        if color is not None:
            run.font.color.rgb = color


def is_table_sep(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return len(cells) >= 1 and all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c != "")


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main(src, dst):
    with io.open(src, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []

    while i < n:
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = MONO_FONT
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # table: a line with '|' followed by a separator row
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(line)
            ncol = len(header)
            body = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                body.append(split_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for c in range(ncol):
                cell_p = hdr[c].paragraphs[0]
                add_inline(cell_p, header[c] if c < len(header) else "", base_bold=True, size=Pt(9))
            for row in body:
                cells = table.add_row().cells
                for c in range(ncol):
                    txt = row[c] if c < len(row) else ""
                    cp = cells[c].paragraphs[0]
                    add_inline(cp, txt, size=Pt(9))
            doc.add_paragraph()
            i = j
            continue

        # headings
        m = re.match(r"(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).replace("**", "")
            if level == 1 and not doc.paragraphs:
                h = doc.add_heading("", level=0)
                add_inline(h, text)
            else:
                h = doc.add_heading("", level=min(level, 4))
                add_inline(h, text)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            qtext = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, qtext)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # bullet list
        if re.match(r"[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
